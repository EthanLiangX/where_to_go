import json
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
from matplotlib import font_manager

import matplotlib.pyplot as plt
import matplotlib.ticker as mtick
from io import BytesIO

# ======================
# 字体设置函数
def set_chinese_font(run, font_name='微软雅黑', font_size=12):
    """正文：微软雅黑，小四"""
    run.font.size = Pt(font_size)
    run.font.name = font_name
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), font_name)

def set_heading_font(paragraph, font_name='微软雅黑', font_size=16):
    """标题：微软雅黑，三号"""
    for run in paragraph.runs:
        run.font.size = Pt(font_size)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# ======================
# matplotlib 设置中文字体（绘图用）
font_path = r"simsun.ttc"  # Windows 下微软雅黑路径
my_font = font_manager.FontProperties(fname=font_path)

# ======================
# 读取 JSON 数据
with open("monitor_data.json", "r", encoding="utf-8") as f:
    monitor_data = json.load(f)

# ======================
# 趋势图生成函数（修复 x/y 维度不一致问题）
def plot_trends(monitor_data, output_dir="charts"):
    os.makedirs(output_dir, exist_ok=True)
    # 收集所有停车场名称，遍历整个数据保证完整
    parking_names = list({d.get("name") for r in monitor_data for d in r.get("data", []) if d.get("name")})
    charts = {}

    for name in parking_names:
        times = []
        usage_rates = []

        for record in monitor_data:
            # 查找当前停车场数据
            p_data = next((d for d in record.get("data", []) if d.get("name") == name), None)
            if p_data:
                times.append(record.get("time"))
                cars = p_data.get("current_cars", 0)
                total = p_data.get("total_slots", 1)
                usage = cars / total if total > 0 else 0
                usage_rates.append(usage)

        # 占用率趋势图
        if usage_rates:
            plt.figure(figsize=(6, 4))
            plt.plot(times, usage_rates, marker='o', label='占用率')
            plt.ylim(0, 1)
            plt.gca().yaxis.set_major_formatter(mtick.PercentFormatter(1.0))  # 转换为百分比格式
            plt.xlabel("时间", fontproperties=my_font)
            plt.ylabel("占用率", fontproperties=my_font)
            plt.title(f"{name} 占用率趋势", fontproperties=my_font)
            plt.grid(True)
            max_idx = usage_rates.index(max(usage_rates))
            min_idx = usage_rates.index(min(usage_rates))
            plt.scatter(times[max_idx], usage_rates[max_idx], color='red', s=100, label='高峰')
            plt.scatter(times[min_idx], usage_rates[min_idx], color='green', s=100, label='低谷')
            plt.legend(prop=my_font)
            plt.tight_layout()
            usage_file = os.path.join(output_dir, f"{name}_usage.png")
            plt.savefig(usage_file)
            plt.close()
        else:
            usage_file = None

        charts[name] = {"usage": usage_file}

    # 所有停车场占用率对比图
    plt.figure(figsize=(8, 5))
    for name in parking_names:
        rates = []
        times_all = []
        for record in monitor_data:
            p_data = next((d for d in record.get("data", []) if d.get("name") == name), None)
            if p_data:
                times_all.append(record.get("time"))
                cars = p_data.get("current_cars", 0)
                total = p_data.get("total_slots", 1)
                usage = cars / total if total > 0 else 0
                rates.append(usage)
        if rates:
            plt.plot(times_all, rates, marker='o', label=name)
    plt.gca().yaxis.set_major_formatter(mtick.PercentFormatter(1.0))  # 转换为百分比格式
    plt.ylim(0, 1)
    plt.xlabel("时间", fontproperties=my_font)
    plt.ylabel("占用率", fontproperties=my_font)

    plt.title("各停车场占用率对比", fontproperties=my_font)
    plt.grid(True)
    plt.legend(prop=my_font)
    plt.tight_layout()
    compare_file = os.path.join(output_dir, "all_parking_compare.png")
    plt.savefig(compare_file)
    plt.close()
    charts["all_compare"] = compare_file

    return charts

# ======================
# 生成单次报告
def generate_single_report(monitor_data):
    current_record = monitor_data[-1]
    charts = plot_trends(monitor_data)
    doc = Document()
    report_time = current_record.get("time", "未知时间")

    # 报告标题
    heading = doc.add_heading("地面停车场无人机监测报告（单次）", 0)
    set_heading_font(heading, font_size=28)

    # # 报告背景
    # heading = doc.add_heading("报告背景", level=1)
    # set_heading_font(heading, font_size=16)
    # p_bg = doc.add_paragraph()
    # set_chinese_font(p_bg.add_run(
    #     "为保障2025年8月在大湾区文化体育中心举办的粤超足球赛交通顺畅与赛事安全，区相关部门自7月起对赛事交通组织进行了系统规划与协调。多轮专题汇报和部门协调会明确了临时停车场建设及交通组织方案，并对方案进行了优化完善。"))
    # p_bg2 = doc.add_paragraph()
    # set_chinese_font(p_bg2.add_run(
    #     "粤超足球赛作为大型体育赛事，将吸引大量观众入场，赛事期间停车需求和周边道路交通压力将显著增加。停车场作为赛事交通核心节点，其使用效率直接关系到观众出行体验及道路通行状况。"))
    # p_bg3 = doc.add_paragraph()
    # set_chinese_font(p_bg3.add_run(
    #     "为实现科学管理，应用无人机对临时停车场及周边道路进行空中监测，可实时获取车位使用率、车辆流量及拥堵情况等关键数据，为临时交通引导、分流策略和管控决策提供数据支撑。"))
    # p_bg4 = doc.add_paragraph()
    # set_chinese_font(p_bg4.add_run(
    #     "本报告基于无人机监测技术，对临时停车场停车情况进行全面分析，旨在评估停车场使用效率、优化赛事期间交通组织方案，并为粤超足球赛的顺利举办提供科学依据。"))
    # # 监测目的
    # heading = doc.add_heading("监测目的", level=1)
    # set_heading_font(heading, font_size=16)
    # items = [
    #     "评估停车场使用效率：实时监测车位占用情况，掌握高峰时段停车压力，为临时停车引导和分配提供数据支撑。",
    #     "分析交通流量与出入口状况：监测停车场主要出入口的车辆进出情况，识别潜在拥堵风险，辅助临时交通组织。",
    #     "支持决策与应急管理：为赛事期间停车调度、交通管控和突发情况处置提供实时数据依据。"
    # ]
    # for item in items:
    #     p = doc.add_paragraph(style="List Bullet")
    #     set_chinese_font(p.add_run(item), font_size=12)

    # 报告基本信息
    heading = doc.add_heading("报告基本信息", level=1)
    set_heading_font(heading, font_size=16)
    info = [
        f"报告编号：UAV-{report_time.replace(':', '')}",
        f"监测时间：{report_time}",
        f"监测停车场数量：{len(current_record.get('data', []))}",
        "数据来源：无人机航拍+AI智能识别",
        "报告类型：单次快报"
    ]
    for line in info:
        p = doc.add_paragraph()
        set_chinese_font(p.add_run(line), font_size=12)

    # 表格
    heading = doc.add_heading("停车场当前状态与历史汇总(现场照片详见附件)", level=1)
    set_heading_font(heading, font_size=16)
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    # headers = ["序号", "停车场名称", "位置", "总车位数", "当前车辆数", "当前占用率", "历史最大占用率","当前总人数"]
    headers = ["序号", "停车场名称", "位置", "总车位数", "当前车辆数", "当前占用率", "历史最大占用率"]
    for idx, text in enumerate(headers):
        run = table.rows[0].cells[idx].paragraphs[0].add_run(text)
        run.bold = True
        set_chinese_font(run, font_size=12)
        table.rows[0].cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 填写数据
    for idx, p_data in enumerate(current_record.get("data", []), start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = p_data.get("name", "未知")
        row_cells[2].text = p_data.get("location", "未知")
        row_cells[3].text = str(p_data.get("total_slots", 0))
        row_cells[4].text = str(p_data.get("current_cars", 0))
        usage_now = p_data.get("current_cars", 0) / p_data.get("total_slots", 1)
        row_cells[5].text = f"{usage_now:.1%}"
        # 历史最大占用率
        max_usage = 0
        for record in monitor_data:
            for d in record.get("data", []):
                if d.get("name") == p_data.get("name"):
                    u = d.get("current_cars", 0) / (d.get("total_slots", 1))
                    max_usage = max(max_usage, u)
        row_cells[6].text = f"{max_usage:.1%}"
        row_cells[7].text = str(p_data.get("current_people",0))
    # 遍历整个表格统一字体
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_chinese_font(run, font_size=12)

    # # 分析与建议 + 图表
    # heading = doc.add_heading("图表分析", level=1)
    # set_heading_font(heading,font_size=16)
    # for p in current_record.get("data", []):
    #     usage_now = p.get("current_cars", 0) / p.get("total_slots", 1)
    #     status = "正常"
    #     if usage_now > 0.9:
    #         status = "接近饱和"
    #     elif usage_now < 0.4:
    #         status = "空闲"
    #
    #     para = doc.add_paragraph(style="List Bullet")
    #     para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #     set_chinese_font(
    #         para.add_run(f"{p.get('name', '未知')}：当前占用率 {usage_now:.1%}，状态 {status}"))
    #
    #     # 判断图表是否存在
    #     chart_data = charts.get(p.get("name"))
    #     para = doc.add_paragraph()
    #     run = para.add_run()
    #     para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #     if chart_data and chart_data["usage"]:
    #         run.add_picture(chart_data["usage"], width=Inches(4))
    #     else:
    #         para_missing = doc.add_paragraph(style="List Bullet")
    #         set_chinese_font(
    #             para_missing.add_run(f"{p.get('name', '未知')}：图表未生成，可能该停车场在监测首条记录缺失。"))

    # 结论与建议
    heading = doc.add_heading("结论与建议", level=1)
    set_heading_font(heading, font_size=16)
    conclusion = [
        "当前停车资源整体充足，但部分区域在高峰时段接近饱和。",
        "建议加强动态调度，优化车流引导。",
        "对高风险时段提前部署疏导人员，确保赛事交通安全。"
    ]
    for line in conclusion:
        p = doc.add_paragraph(style="List Bullet")
        set_chinese_font(p.add_run(line), font_size=12)
    heading = doc.add_heading("附件", level=1)
    set_heading_font(heading, font_size=16)
    filename = f"大湾区文化体育中心粤超足球赛活动停车场无人机监测单次报告_{report_time.replace(':', '点')}分.docx"
    doc.save(filename)
    print(f"✅ 已生成单次报告：{filename}")

# ======================
# 生成全天汇总报告
def generate_full_day_report(monitor_data):
    charts = plot_trends(monitor_data)
    doc = Document()

    heading = doc.add_heading("地面停车场无人机监测报告（全天汇总）", 0)
    set_heading_font(heading, font_size=28)

    # 报告背景
    heading = doc.add_heading("报告背景", level=1)
    set_heading_font(heading, font_size=16)
    p_bg = doc.add_paragraph()
    set_chinese_font(p_bg.add_run(
        "为保障2025年8月在大湾区文化体育中心举办的粤超足球赛交通顺畅与赛事安全，区相关部门自7月起对赛事交通组织进行了系统规划与协调。多轮专题汇报和部门协调会明确了临时停车场建设及交通组织方案，并对方案进行了优化完善。"))
    p_bg2 = doc.add_paragraph()
    set_chinese_font(p_bg2.add_run(
        "粤超足球赛作为大型体育赛事，将吸引大量观众入场，赛事期间停车需求和周边道路交通压力将显著增加。停车场作为赛事交通核心节点，其使用效率直接关系到观众出行体验及道路通行状况。"))
    p_bg3 = doc.add_paragraph()
    set_chinese_font(p_bg3.add_run(
        "为实现科学管理，应用无人机对临时停车场及周边道路进行空中监测，可实时获取车位使用率、车辆流量及拥堵情况等关键数据，为临时交通引导、分流策略和管控决策提供数据支撑。"))
    p_bg4 = doc.add_paragraph()
    set_chinese_font(p_bg4.add_run(
        "本报告基于无人机监测技术，对临时停车场停车情况进行全面分析，旨在评估停车场使用效率、优化赛事期间交通组织方案，并为粤超足球赛的顺利举办提供科学依据。"))
    # 监测目的
    heading = doc.add_heading("监测目的", level=1)
    set_heading_font(heading, font_size=16)
    p_purpose1 = doc.add_paragraph(style="List Bullet")
    set_chinese_font(p_purpose1.add_run(
        "评估停车场使用效率：实时监测车位占用情况，掌握高峰时段停车压力，为临时停车引导和分配提供数据支撑。"))
    p_purpose2 = doc.add_paragraph(style="List Bullet")
    set_chinese_font(p_purpose2.add_run(
        "分析交通流量与出入口状况：监测停车场主要出入口的车辆进出情况，识别潜在拥堵风险，辅助临时交通组织。"))
    p_purpose3 = doc.add_paragraph(style="List Bullet")
    set_chinese_font(
        p_purpose3.add_run("支持决策与应急管理：为赛事期间停车调度、交通管控和突发情况处置提供实时数据依据。"))
    info = [
        f"报告编号：UAV-FULLDAY",
        f"监测日期/周期：{monitor_data[0].get('time', '未知')} 至 {monitor_data[-1].get('time', '未知')}",
        f"监测停车场数量：{len(monitor_data[0].get('data', []))}",
        "数据来源：无人机航拍+AI智能识别",
        "报告类型：全天监测总结报告"
    ]
    for line in info:
        p = doc.add_paragraph()
        set_chinese_font(p.add_run(line), font_size=12)
    # 数据汇总表
    doc.add_heading("全天停车场数据汇总", level=1)
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    headers = ["序号", "停车场名称", "位置", "总车位数", "平均占用率", "最大占用率", "最小占用率", "总监测次数"]
    for idx, text in enumerate(headers):
        run = table.rows[0].cells[idx].paragraphs[0].add_run(text)
        run.bold = True
        set_chinese_font(run, font_size=12)
        table.rows[0].cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 收集所有停车场名称
    parking_names = [d.get("name") for d in monitor_data[0].get("data", []) if d.get("name")]
    for idx, name in enumerate(parking_names, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        first_p = next((d for r in monitor_data for d in r.get("data", []) if d.get("name") == name), {})
        row_cells[1].text = first_p.get("name", "未知")
        row_cells[2].text = first_p.get("location", "未知")
        row_cells[3].text = str(first_p.get("total_slots", 0))
        usage_list = [
            d.get("current_cars", 0) / d.get("total_slots", 1)
            for r in monitor_data
            for d in r.get("data", [])
            if d.get("name") == name
        ]

        row_cells[4].text = f"{sum(usage_list) / len(usage_list):.1%}" if usage_list else "0%"
        row_cells[5].text = f"{max(usage_list):.1%}" if usage_list else "0%"
        row_cells[6].text = f"{min(usage_list):.1%}" if usage_list else "0%"
        row_cells[7].text = str(len(usage_list))
    # 分析与建议 + 图表
    heading = doc.add_heading("图表分析", level=1)
    set_heading_font(heading, font_size=16)
    for name in parking_names:
        usage_list = [
            d.get("current_cars", 0) / d.get("total_slots", 1)
            for r in monitor_data
            for d in r.get("data", [])
            if d.get("name") == name
        ]
        if not usage_list: continue
        avg_usage = sum(usage_list) / len(usage_list)
        max_idx = usage_list.index(max(usage_list))
        min_idx = usage_list.index(min(usage_list))
        high_time = monitor_data[max_idx]["time"]
        low_time = monitor_data[min_idx]["time"]
        status = "正常"
        if avg_usage > 0.9:
            status = "全天高负荷"
        elif avg_usage < 0.4:
            status = "全天空闲"
        para = doc.add_paragraph(style="List Bullet")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_chinese_font(
            para.add_run(
                f"{name}：平均占用率 {avg_usage:.1%}，状态 {status}，高峰时间 {high_time}，低谷时间 {low_time}"))

        chart_data = charts.get(name)
        para = doc.add_paragraph()
        run = para.add_run()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if chart_data and chart_data["usage"]:
            run.add_picture(chart_data["usage"], width=Inches(4))
        else:
            para_missing = doc.add_paragraph(style="List Bullet")
            set_chinese_font(para_missing.add_run(f"{name}：图表未生成，可能该停车场在监测首条记录缺失。"))

    # 所有停车场占用率对比图
    heading = doc.add_heading("各停车场占用率对比", level=1)
    set_heading_font(heading, font_size=16)
    doc.add_picture(charts["all_compare"], width=Inches(6))
    heading = doc.add_heading("总体与建议", level=1)
    set_heading_font(heading, font_size=16)
    summary = [
        "全天停车场总体运行平稳，绝大多数时段未出现明显饱和。",
        "个别停车场在比赛前2小时出现高峰，接近饱和。",
        "建议建立动态车流预测模型，提前调配交通资源。"
    ]
    for line in summary:
        p = doc.add_paragraph(style="List Bullet")
        set_chinese_font(p.add_run(line), font_size=12)

    filename = "大湾区文化体育中心粤超足球赛活动停车场无人机监测全天汇总报告.docx"
    doc.save(filename)
    print(f"✅ 已生成全天汇总报告：{filename}")

# ======================
# 执行
generate_single_report(monitor_data)
generate_full_day_report(monitor_data)
